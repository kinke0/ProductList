import subprocess
import sys
import os
import urllib.parse

def check_and_install_dependencies():
    required_packages = {
        'pandas': 'pandas',
        'openpyxl': 'openpyxl',
        'docx': 'python-docx',
        'requests': 'requests',
        'PIL': 'Pillow'
    }

check_and_install_dependencies()

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import requests
from io import BytesIO
from PIL import Image

def parse_arguments():
    if len(sys.argv) < 2:
        print("用法: python generate_word.py <输入Excel文件路径> [输出Word文件路径]")
        print("  输入Excel文件路径: 必需的，指定要读取的Excel文件路径")
        print("  输出Word文件路径: 可选的，指定要生成的Word文件路径")
        print("                   如果不指定，则使用与输入文件相同的名称（扩展名为.docx）")
        sys.exit(1)

    input_file = sys.argv[1]

    if not os.path.isfile(input_file):
        print(f"错误: 输入文件不存在 '{input_file}'")
        sys.exit(1)

    if len(sys.argv) >= 3:
        output_file = sys.argv[2]
    else:
        base_name = os.path.splitext(input_file)[0]
        output_file = base_name + '.docx'

    return input_file, output_file

input_file, output_file = parse_arguments()
print(f"输入文件: {input_file}")
print(f"输出文件: {output_file}")

df = pd.read_excel(input_file)
required_columns = ['产品/系统', '父记录', '功能说明', '业务分类', '业务域']
df = df[required_columns]
df = df.dropna(how='all')

doc = Document()
doc.styles['Normal'].paragraph_format.line_spacing = 1.5

def set_font_style(run):
    run.font.name = '宋体'
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.italic = False
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading_with_number(doc: Document, text: str, level: int, number: str):
    """
    使用原生Heading样式添加带编号的标题

    Args:
        doc: Word文档对象
        text: 标题文本内容
        level: 标题级别，范围1-9
        number: 编号字符串，如"1", "1.1", "1.1.1"

    Returns:
        生成的段落对象
    """
    full_text = f"{number} {text}"
    
    para = doc.add_heading(full_text, level=level)
    
    for run in para.runs:
        run.font.name = '宋体'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.italic = False
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    return para

def fix_image_url(url):
    """修复图片URL中的常见问题"""
    import urllib.parse
    
    # 1. 修复缺少扩展名分隔符的情况，如 "27png" -> "27.png"
    url = re.sub(r'(\d)(png|jpg|jpeg|gif|bmp|webp)$', r'\1.\2', url, flags=re.IGNORECASE)
    
    # 2. 对URL进行编码，处理空格和特殊字符
    # 先解析URL
    parsed = urllib.parse.urlparse(url)
    # 对路径进行编码，保留斜杠和已编码的字符（如%20）
    # 先解码已编码的部分，避免双重编码
    decoded_path = urllib.parse.unquote(parsed.path)
    encoded_path = urllib.parse.quote(decoded_path, safe='/')
    # 重新组装URL
    fixed_url = urllib.parse.urlunparse((
        parsed.scheme,
        parsed.netloc,
        encoded_path,
        parsed.params,
        parsed.query,
        parsed.fragment
    ))
    
    return fixed_url

def download_image(url, max_width=6.0):
    try:
        if os.path.isabs(url) and os.path.isfile(url):
            if os.path.exists(url):
                return open(url, 'rb'), None
            else:
                return None, f"本地文件不存在: {url}"
        
        # 修复URL
        original_url = url
        url = fix_image_url(url)
        if url != original_url:
            print(f"  URL已修复: {url}")
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'image/webp,image/apng,image/*,*/*;q=0.8',
            'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
            'Referer': 'https://feishu.cn/'
        }
        response = requests.get(url, headers=headers, timeout=15, allow_redirects=True, stream=True)
        content_type = response.headers.get('Content-Type', '')
        content_length = response.headers.get('Content-Length', 'unknown')
        print(f"  Content-Type: {content_type}, Content-Length: {content_length}")
        if response.status_code == 200:
            if 'image' in content_type or any(ext in url.lower() for ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp']):
                image_data = BytesIO(response.content)
                return image_data, None
            elif 'text' in content_type or response.status_code == 401:
                return None, f"需要登录认证才能访问"
            else:
                return None, f"URL不是图片格式 (Content-Type: {content_type})"
        else:
            return None, f"下载失败: HTTP {response.status_code}"
    except requests.exceptions.Timeout:
        return None, "下载失败: 请求超时"
    except requests.exceptions.SSLError as e:
        return None, f"下载失败: SSL证书错误"
    except Exception as e:
        return None, f"下载失败: {str(e)}"

def convert_image_to_png(image_data):
    try:
        image_data.seek(0)
        img = Image.open(image_data)
        if img.format == 'WEBP':
            print(f"  检测到WebP格式，转换为PNG...")
            output = BytesIO()
            img.save(output, format='PNG')
            output.seek(0)
            return output, None
        return image_data, None
    except Exception as e:
        return None, f"图片格式转换失败: {str(e)}"

def get_image_dimensions(image_data):
    try:
        image_data.seek(0)
        img = Image.open(image_data)
        width, height = img.size
        return width, height
    except Exception as e:
        return None, None

def insert_single_image(doc, url, image_data, height_px):
    width, height = get_image_dimensions(image_data)
    if width and height:
        if height > width * 1.2:
            target_height = 300
            aspect_ratio = width / height
            target_width = target_height * aspect_ratio
        else:
            target_width = 500
            aspect_ratio = height / width if width > 0 else 1
            target_height = target_width * aspect_ratio
    else:
        target_width = 600
        target_height = 250
    
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(21)
    run = para.add_run()
    run.add_picture(image_data, height=Inches(target_height / 96), width=Inches(target_width / 96))
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parsed = urllib.parse.urlparse(url)
    decoded_path = urllib.parse.unquote(parsed.path)
    filename = os.path.basename(decoded_path).split('?')[0]
    filename_without_ext = os.path.splitext(filename)[0]
    caption_para = doc.add_paragraph()
    caption_run = caption_para.add_run("图：" + filename_without_ext)
    caption_run.font.name = '宋体'
    caption_run.font.size = Pt(10)
    caption_run.font.color.rgb = RGBColor(128, 128, 128)
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def process_description_with_images(doc, description):
    url_pattern = r'<([^<>]+)>'
    parts = re.split(url_pattern, description)
    
    i = 0
    while i < len(parts):
        part = parts[i].strip()
        if not part:
            i += 1
            continue
        
        if re.match(r'^https?://', part):
            consecutive_urls = [part]
            j = i + 1
            while j < len(parts):
                next_part = parts[j].strip()
                if not next_part:
                    j += 1
                    continue
                if re.match(r'^https?://', next_part):
                    consecutive_urls.append(next_part)
                    j += 1
                else:
                    break
            
            if len(consecutive_urls) >= 2:
                is_vertical_group = False
                group_images_data = []
                for url in consecutive_urls:
                    image_data, error = download_image(url)
                    if error:
                        break
                    image_data, convert_error = convert_image_to_png(image_data)
                    if convert_error:
                        break
                    image_data.seek(0)
                    width, height = get_image_dimensions(image_data)
                    if width and height:
                        if height > width * 1.2:
                            is_vertical_group = True
                        image_data.seek(0)
                        group_images_data.append((url, image_data, width, height))
                
                num_cols = min(len(group_images_data), 3)
                if is_vertical_group and len(group_images_data) >= 2:
                    print(f"  检测到连续竖版截图 {len(group_images_data)} 张，进行每行{num_cols}张排列")
                    
                    rows_data = []
                    current_row = []
                    for idx, (url, image_data, width, height) in enumerate(group_images_data):
                        current_row.append((url, image_data, width, height))
                        if len(current_row) == num_cols or idx == len(group_images_data) - 1:
                            rows_data.append(current_row)
                            current_row = []
                    
                    table = doc.add_table(rows=len(rows_data), cols=num_cols)
                    
                    tblPr = table._element.find(qn('w:tblPr'))
                    if tblPr is None:
                        tblPr = OxmlElement('w:tblPr')
                        table._element.insert(0, tblPr)
                    tblStyle = tblPr.find(qn('w:tblStyle'))
                    if tblStyle is not None:
                        tblPr.remove(tblStyle)
                    existing_borders = tblPr.find(qn('w:tblBorders'))
                    if existing_borders is not None:
                        tblPr.remove(existing_borders)
                    tblBorders = OxmlElement('w:tblBorders')
                    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                        border = OxmlElement(f'w:{border_name}')
                        border.set(qn('w:val'), 'single')
                        border.set(qn('w:sz'), '4')
                        border.set(qn('w:color'), 'FFFFFF')
                        border.set(qn('w:space'), '0')
                        tblBorders.append(border)
                    tblPr.append(tblBorders)
                    
                    for ri, row_urls in enumerate(rows_data):
                        row_cells = table.rows[ri].cells
                        for ci, (url, image_data, width, height) in enumerate(row_urls):
                            cell = row_cells[ci]
                            tcPr = cell._element.find(qn('w:tcPr'))
                            if tcPr is None:
                                tcPr = OxmlElement('w:tcPr')
                                cell._element.insert(0, tcPr)
                            existing_tc_bdr = tcPr.find(qn('w:tcBdr'))
                            if existing_tc_bdr is not None:
                                tcPr.remove(existing_tc_bdr)
                            
                            target_height = 300
                            aspect_ratio = width / height if height > 0 else 1
                            target_width = target_height * aspect_ratio
                            
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = cell.paragraphs[0].add_run()
                            run.add_picture(image_data, height=Inches(target_height / 96), width=Inches(target_width / 96))
                            
                            caption_para = cell.add_paragraph()
                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            parsed = urllib.parse.urlparse(url)
                            decoded_path = urllib.parse.unquote(parsed.path)
                            filename = os.path.basename(decoded_path).split('?')[0]
                            filename_without_ext = os.path.splitext(filename)[0]
                            caption_run = caption_para.add_run("图：" + filename_without_ext)
                            caption_run.font.name = '宋体'
                            caption_run.font.size = Pt(8)
                            caption_run.font.color.rgb = RGBColor(128, 128, 128)
                        
                        if len(row_urls) < num_cols:
                            for empty_ci in range(len(row_urls), num_cols):
                                row_cells[empty_ci].text = ''
                    
                    print(f"  连续竖版截图排列完成")
                    i = j
                    continue
                else:
                    for url in consecutive_urls:
                        print(f"  正在下载图片: {url[:80]}...")
                        image_data, error = download_image(url)
                        print(f"  下载结果: error={error}, image_data={type(image_data)}")
                        if error:
                            print(f"  {error}，尝试加载错误提示图片")
                            error_image_url = 'http://cloudimgs.jscloud.vip:16666/api/images/%E6%8F%90%E7%A4%BA%E5%9B%BE%E7%89%87/error.png'
                            image_data, error2 = download_image(error_image_url)
                            if error2:
                                print(f"  错误提示图片也加载失败: {error2}，保留原文")
                                para = doc.add_paragraph()
                                para.paragraph_format.first_line_indent = Pt(21)
                                run = para.add_run(f"<{url}>")
                                set_font_style(run)
                                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            else:
                                try:
                                    image_data, convert_error = convert_image_to_png(image_data)
                                    if convert_error:
                                        print(f"  {convert_error}，保留原文")
                                        para = doc.add_paragraph()
                                        para.paragraph_format.first_line_indent = Pt(21)
                                        run = para.add_run(f"<{url}>")
                                        set_font_style(run)
                                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                    else:
                                        insert_single_image(doc, url, image_data, None)
                                        print(f"  错误提示图片插入成功")
                                except Exception as e:
                                    import traceback
                                    print(f"  错误提示图片插入失败: {str(e)}")
                                    para = doc.add_paragraph()
                                    para.paragraph_format.first_line_indent = Pt(21)
                                    run = para.add_run(f"<{url}>")
                                    set_font_style(run)
                                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        else:
                            try:
                                image_data, convert_error = convert_image_to_png(image_data)
                                if convert_error:
                                    print(f"  {convert_error}，保留原文")
                                    para = doc.add_paragraph()
                                    para.paragraph_format.first_line_indent = Pt(21)
                                    run = para.add_run(f"<{url}>")
                                    set_font_style(run)
                                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                else:
                                    image_data.seek(0)
                                    image_size = len(image_data.getvalue())
                                    print(f"  图片大小: {image_size} bytes")
                                    if image_size == 0:
                                        print(f"  图片数据为空，保留原文")
                                        para = doc.add_paragraph()
                                        para.paragraph_format.first_line_indent = Pt(21)
                                        run = para.add_run(f"<{url}>")
                                        set_font_style(run)
                                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                    else:
                                        insert_single_image(doc, url, image_data, None)
                                        print(f"  图片插入成功")
                            except Exception as e:
                                import traceback
                                print(f"  图片插入失败: {str(e)}")
                                para = doc.add_paragraph()
                                para.paragraph_format.first_line_indent = Pt(21)
                                run = para.add_run(f"<{url}>")
                                set_font_style(run)
                                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    i = j
                    continue
            else:
                url = consecutive_urls[0]
                print(f"  正在下载图片: {url[:80]}...")
                image_data, error = download_image(url)
                print(f"  下载结果: error={error}, image_data={type(image_data)}")
                if error:
                    print(f"  {error}，尝试加载错误提示图片")
                    error_image_url = 'http://cloudimgs.jscloud.vip:16666/api/images/%E6%8F%90%E7%A4%BA%E5%9B%BE%E7%89%87/error.png'
                    image_data, error2 = download_image(error_image_url)
                    if error2:
                        print(f"  错误提示图片也加载失败: {error2}，保留原文")
                        para = doc.add_paragraph()
                        para.paragraph_format.first_line_indent = Pt(21)
                        run = para.add_run(f"<{url}>")
                        set_font_style(run)
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        try:
                            image_data, convert_error = convert_image_to_png(image_data)
                            if convert_error:
                                print(f"  {convert_error}，保留原文")
                                para = doc.add_paragraph()
                                para.paragraph_format.first_line_indent = Pt(21)
                                run = para.add_run(f"<{url}>")
                                set_font_style(run)
                                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            else:
                                insert_single_image(doc, url, image_data, None)
                                print(f"  错误提示图片插入成功")
                        except Exception as e:
                            import traceback
                            print(f"  错误提示图片插入失败: {str(e)}")
                            para = doc.add_paragraph()
                            para.paragraph_format.first_line_indent = Pt(21)
                            run = para.add_run(f"<{url}>")
                            set_font_style(run)
                            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    try:
                        image_data, convert_error = convert_image_to_png(image_data)
                        if convert_error:
                            print(f"  {convert_error}，保留原文")
                            para = doc.add_paragraph()
                            para.paragraph_format.first_line_indent = Pt(21)
                            run = para.add_run(f"<{url}>")
                            set_font_style(run)
                            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        else:
                            image_data.seek(0)
                            image_size = len(image_data.getvalue())
                            print(f"  图片大小: {image_size} bytes")
                            if image_size == 0:
                                print(f"  图片数据为空，保留原文")
                                para = doc.add_paragraph()
                                para.paragraph_format.first_line_indent = Pt(21)
                                run = para.add_run(f"<{url}>")
                                set_font_style(run)
                                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            else:
                                insert_single_image(doc, url, image_data, None)
                                print(f"  图片插入成功")
                    except Exception as e:
                        import traceback
                        print(f"  图片插入失败: {str(e)}")
                        para = doc.add_paragraph()
                        para.paragraph_format.first_line_indent = Pt(21)
                        run = para.add_run(f"<{url}>")
                        set_font_style(run)
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                i = j
                continue
        else:
            lines = part.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    para = doc.add_paragraph()
                    para.paragraph_format.first_line_indent = Pt(21)
                    run = para.add_run(line)
                    set_font_style(run)
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue
        
        i += 1



def extract_text(code_text):
    text = str(code_text).strip()
    if '@' in text:
        text = text.split('@')[0].strip()
    text = re.sub(r'^[\d\.]+\s+', '', text)
    return text.strip()

def extract_code(product):
    if isinstance(product, str):
        text = str(product).strip().replace('\n', '').replace('\r', '')
        match = re.match(r'^[\d\.]+', text)
        if match:
            return match.group(0)
    return None

def extract_name(product):
    if isinstance(product, str):
        text = str(product).strip().replace('\n', '').replace('\r', '')
        match = re.match(r'^[\d\.]+\s*(.*)', text)
        if match:
            return match.group(1).strip()
    return product

def _get_code_sort_key(code: str):
    """
    获取节点代码的排序键
    
    Args:
        code: 节点代码（如 "1", "1.1", "1.1.1"）
    
    Returns:
        用于排序的元组，如 (1,), (1, 1), (1, 1, 1)
    """
    return tuple(int(p) for p in code.split('.'))


def build_numbering_map(nodes: dict, children_map: dict) -> dict:
    """
    构建节点编号映射（1, 1.1, 1.1.1...）
    
    Args:
        nodes: 节点字典 {code: {name, parent_code, ...}}
        children_map: 父节点到子节点列表的映射 {parent_code: [child_code, ...]}
    
    Returns:
        编号映射 {code: "1", "1.1", "1.1.1", ...}
    
    Raises:
        ValueError: 当nodes为空或None时抛出
    """
    if not nodes:
        raise ValueError("nodes必须包含至少一个节点")
    
    numbering_map = {}
    
    l1_codes = [code for code in nodes if '.' not in code]
    l1_codes.sort(key=lambda x: int(x) if x.isdigit() else 0)
    
    for i, code in enumerate(l1_codes, 1):
        numbering_map[code] = str(i)
        if children_map is not None:
            process_children(code, nodes, children_map, numbering_map, str(i))
    
    return numbering_map

def process_children(parent_code: str, nodes: dict, children_map: dict, 
                     numbering_map: dict, parent_num: str) -> None:
    """
    递归处理子节点的编号
    
    Args:
        parent_code: 父节点代码
        nodes: 节点字典
        children_map: 子节点映射
        numbering_map: 编号映射（会被修改）
        parent_num: 父节点编号
    
    Returns:
        None
    """
    if not parent_code or not children_map or not numbering_map:
        return
    
    if parent_code not in children_map:
        return
    
    children = children_map[parent_code]
    children.sort(key=_get_code_sort_key)
    
    for i, child_code in enumerate(children, 1):
        child_num = f"{parent_num}.{i}"
        numbering_map[child_code] = child_num
        process_children(child_code, nodes, children_map, numbering_map, child_num)

def process_data():
    unique_categories = df['业务分类'].dropna().unique()
    category_counter = 0
    
    for category in unique_categories:
        if isinstance(category, str):
            category_counter += 1
            category_text = extract_text(category)
            category_number = str(category_counter)

            category_para = add_heading_with_number(doc, category_text, 1, category_number)

            category_df = df[df['业务分类'] == category].copy()
            
            unique_domains = []
            seen_domains = set()
            for domain in category_df['业务域'].dropna():
                if domain not in seen_domains:
                    unique_domains.append(domain)
                    seen_domains.add(domain)
            
            domain_counter = 0
            
            for domain in unique_domains:
                if isinstance(domain, str):
                    domain_counter += 1
                    domain_text = extract_text(domain)
                    domain_number = f"{category_number}.{domain_counter}"

                    domain_para = add_heading_with_number(doc, domain_text, 2, domain_number)

                    domain_df = category_df[category_df['业务域'] == domain].copy()
                    
                    nodes = {}
                    for _, row in domain_df.iterrows():
                        product = row['产品/系统']
                        parent_record = row['父记录']
                        code = extract_code(product)
                        
                        if code:
                            parent_code = None
                            if isinstance(parent_record, str) and parent_record.strip():
                                parent_code = extract_code(parent_record)
                            
                            nodes[code] = {
                                'name': extract_name(product),
                                'parent_code': parent_code,
                                'description': row['功能说明']
                            }
                    
                    children_map = {}
                    for code in nodes:
                        parent_code = nodes[code]['parent_code']
                        if parent_code and parent_code in nodes:
                            if parent_code not in children_map:
                                children_map[parent_code] = []
                            children_map[parent_code].append(code)
                    
                    numbering_map = build_numbering_map(nodes, children_map)
                    
                    processed_nodes = set()
                    
                    def process_node_with_numbering(code, parent_number=None):
                        if code in processed_nodes or code not in nodes:
                            return
                        
                        current_number = numbering_map.get(code, '')
                        actual_doc_level = len(current_number.split('.')) + 2
                        actual_doc_level = min(max(actual_doc_level, 3), 9)
                        
                        name = nodes[code]['name']
                        product_para = add_heading_with_number(doc, name, actual_doc_level, current_number)
                        
                        description = nodes[code]['description']
                        if isinstance(description, str) and description.strip():
                            process_description_with_images(doc, description)
                        
                        if code in children_map:
                            for child_code in children_map[code]:
                                process_node_with_numbering(child_code, current_number)
                    
                    level_3_codes = [code for code in nodes if len(code.split('.')) == 3]
                    level_3_codes.sort(key=lambda x: tuple(int(p) for p in x.split('.')))
                    
                    for code in level_3_codes:
                        process_node_with_numbering(code)

process_data()
doc.save(output_file)
print(f"文档生成完成：{output_file}")
